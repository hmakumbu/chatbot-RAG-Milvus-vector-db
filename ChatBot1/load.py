from langchain_community.document_loaders import TextLoader
from docx import Document
from pypdf import PdfReader
from langchain.schema import Document
#import fitz  # PyMuPDF for PDF handling
import os


def load_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    loader = TextLoader(text)
    documents = loader.load()
    return documents

def load_text(file_path):
   
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"The file {file_path} does not exist.")
    
    loader = TextLoader(file_path)
    documents = loader.load()
    
    #print(documents)
    #print(type(documents))
    return documents

# def load_pdf(file_path):
#     # # pdf_document = fitz.open(file_path)
#     # # text = ""
#     # # for page_num in range(len(pdf_document)):
#     # #     page = pdf_document.load_page(page_num)
#     # #     text += page.get_text()

#     # # loader = TextLoader(text)
#     # # documents = loader.load()
#     # reader = PdfReader(file_path)
#     # pdf_texts = [p.extract_text().strip() for p in reader.pages]
#     # # Filter the empty strings
#     # documents = [text for text in pdf_texts if text]
#     # # return data
#     if not os.path.exists(file_path):
#         raise FileNotFoundError(f"The file {file_path} does not exist.")
    
#     reader = PdfReader(file_path)
#     pdf_texts = [page.extract_text().strip() for page in reader.pages]
#     documents = [text for text in pdf_texts if text]
    
#     return documents


def load_pdf(file_path):

    try:
        reader = PdfReader(file_path)
        pdf_texts = [p.extract_text().strip() for p in reader.pages if p is not None]
        documents = []
        for text in pdf_texts:
            metadata = {'source': file_path}
            doc = Document(metadata=metadata, page_content=text)
            documents.append(doc)
            print(documents)
            print(type(documents))
        return documents
    except Exception as e:
        print(f"Error loading PDF file '{file_path}': {str(e)}")
        raise RuntimeError(f"Error loading {file_path}") from e

def load(file_path):

    file_extension = file_path.split(".")[-1].lower()

    if file_extension == "txt":
      
        return load_text(file_path)
    elif file_extension == "pdf":
        return load_pdf(file_path)
    elif file_extension == "docx":
        return load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {file_extension}")













# ------------------------------------------------------------------
  # Read pdf
    # reader = PdfReader(file_path)
    # pdf_texts = [p.extract_text().strip() for p in reader.pages]
    # # Filter the empty strings
    # data = [text for text in pdf_texts if text]
    # return data
# ------------------------------------------------------------------



# from langchain_community.document_loaders import TextLoader
# import os
# from dotenv import load_dotenv

# load_dotenv()

# def load_text(file_path):
#     if not os.path.exists(file_path):
#         raise FileNotFoundError(f"The file {file_path} does not exist.")

#     loader = TextLoader(file_path)
#     documents = loader.load()

#     return documents


